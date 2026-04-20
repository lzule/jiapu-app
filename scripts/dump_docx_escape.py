from pathlib import Path
from docx import Document

folder = Path(r"C:\Users\20163\Desktop\jiapu_app")
p = next(x for x in folder.iterdir() if x.suffix.lower() == ".docx")
d = Document(str(p))
print("file:", p.name, "paragraphs:", len(d.paragraphs))
for i, para in enumerate(d.paragraphs):
    t = para.text.strip()
    if not t:
        continue
    print(f"{i:03d}\t{para.style.name}\t{t.encode('unicode_escape').decode('ascii')}")
