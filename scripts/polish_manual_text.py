# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document


def main():
    folder = Path(r"C:\Users\20163\Desktop\jiapu_app")
    docx_files = [x for x in folder.iterdir() if x.suffix.lower() == ".docx"]
    if not docx_files:
        raise SystemExit("NO_DOCX")

    p = docx_files[0]
    d = Document(str(p))

    replacements = {
        12: "节点上鼠标左键双击：编辑该节点信息。",
        16: "如何建立节点关系：",
        17: "方式一：按住 Alt 并将一个节点拖拽到另一个节点，可快速建立关系。",
        18: "随后在弹出的关系菜单中，可选择：配偶、父母、子女、兄弟姐妹。",
        19: "关系建立后，画面会自动更新，无需手动刷新。",
        20: "方式二：右键目标节点，在菜单中直接创建与其相关的新节点。",
        21: "新创建的节点会根据关系自动排布。",
        41: "同一父母下的子女可通过鼠标拖拽进行排序。",
        43: "1. 可先在空白处鼠标左键双击，快速创建若干独立节点，再通过 Alt + 鼠标拖拽建立节点关系。",
        44: "2. 也可先创建一个节点，再右键该节点，逐步创建与其直接相关的其他节点。",
        45: "3. 合理使用快捷键，可明显提升节点创建和编辑效率。",
        46: "4. 右键菜单或 Alt + 拖拽后的关系菜单中，带字母的选项支持按键快速触发。",
        47: "六、注意事项（避免踩坑）",
        53: "七、常见问题",
    }

    for idx, txt in replacements.items():
        if idx < len(d.paragraphs):
            d.paragraphs[idx].text = txt

    out = folder / "operation_manual_polished_temp.docx"
    d.save(str(out))
    print(out)


if __name__ == "__main__":
    main()

