# -*- coding: utf-8 -*-
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def set_font(style, name, size):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_meta(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].italic = True


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_manual(output_path: Path):
    doc = Document()

    set_font(doc.styles["Normal"], "微软雅黑", 11)
    set_font(doc.styles["Heading 1"], "微软雅黑", 16)
    set_font(doc.styles["Heading 2"], "微软雅黑", 13)
    set_font(doc.styles["Heading 3"], "微软雅黑", 12)

    add_title(doc, "家谱管理离线网页《操作手册》")
    add_meta(doc, "版本：当前稳定版")
    add_meta(doc, "生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    doc.add_paragraph(
        "这是一份快速上手手册，重点告诉你“怎么立刻用起来”。"
        "不讲复杂细节，按下面做即可。"
    )

    doc.add_paragraph("一、3分钟快速上手", style="Heading 1")
    add_numbers(
        doc,
        [
            "用 Chrome 或 Edge 打开 jiapu.html。",
            "点击顶部“连接”，选择你的数据目录。",
            "先添加1位成员，再右键该成员添加配偶/子女/父母/兄弟姐妹。",
            "完成一小段录入后，点击“保存版本”。",
        ],
    )

    doc.add_paragraph("二、最常用操作（必看）", style="Heading 1")
    doc.add_paragraph("1）节点相关", style="Heading 2")
    add_bullets(
        doc,
        [
            "空白处鼠标左键双击：快速创建独立节点。",
            "节点上鼠标左键双击：编辑该节点信息。",
            "选中节点后按 Del：删除该节点。",
            "节点上右键：打开关系操作菜单。",
        ],
    )

    doc.add_paragraph("2）关系相关", style="Heading 2")
    add_bullets(
        doc,
        [
            "Alt + 拖拽一个节点到另一个节点：快速建立关系。",
            "弹出的关系菜单可选：配偶、父母、子女、兄弟姐妹。",
            "关系建立后画面会自动更新，无需手动刷新。",
        ],
    )

    doc.add_paragraph("3）查看与排版", style="Heading 2")
    add_bullets(
        doc,
        [
            "按住鼠标左键拖动画布：平移查看。",
            "滚轮：缩放。",
            "点击“适应”：快速回到全局视图。",
            "拖拽兄弟姐妹节点：调整同辈顺序。",
        ],
    )

    doc.add_paragraph("三、快捷方式一览", style="Heading 1")
    add_bullets(
        doc,
        [
            "Ctrl + Z：撤销",
            "Ctrl + Y：前进",
            "Ctrl + S：保存版本",
            "Ctrl + F：搜索",
            "Ctrl + B：切换侧边栏",
            "Ctrl + /：快捷键设置",
            "Delete：删除当前选中节点",
            "Esc：关闭弹窗/取消当前操作",
        ],
    )

    doc.add_paragraph("四、你最需要知道的规则", style="Heading 1")
    add_bullets(
        doc,
        [
            "单亲时，子女线从该父/母节点引出。",
            "双亲时，子女线从父母中间引出。",
            "从单亲变双亲后，连线会自动切换。",
            "同一父母下的子女顺序会保持一致。",
        ],
    )

    doc.add_paragraph("五、注意事项（避免踩坑）", style="Heading 1")
    add_bullets(
        doc,
        [
            "建议一直用同一个数据目录，不要来回切换。",
            "删除节点前先保存版本，防止误删。",
            "做大改前后各保存一次版本，方便回退。",
            "尽量统一姓名和日期写法，后续搜索更准。",
            "不建议手改数据文件，优先用界面操作。",
        ],
    )

    doc.add_paragraph("六、常见问题", style="Heading 1")
    doc.add_paragraph("Q：为什么关系线位置会变化？", style="Heading 2")
    doc.add_paragraph("A：当家庭从单亲变成双亲时，系统会自动改为从父母中间引线，这是正常行为。")
    doc.add_paragraph("Q：为什么拖拽一个孩子会影响同组顺序？", style="Heading 2")
    doc.add_paragraph("A：系统会统一同一父母组下的子女顺序，避免左右不一致。")
    doc.add_paragraph("Q：数据如何防丢？", style="Heading 2")
    doc.add_paragraph("A：定期备份 data 目录；关键阶段保存版本。")

    doc.save(str(output_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    output = root / "操作手册.docx"
    build_manual(output)
    print(str(output))
